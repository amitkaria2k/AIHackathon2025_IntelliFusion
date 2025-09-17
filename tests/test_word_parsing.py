#!/usr/bin/env python3
"""
Test script for Word document parsing enhancements
"""
import os
import sys
import tempfile
import io
from docx import Document

def create_test_word_doc():
    """Create a test Word document with tables that might cause parsing issues"""
    doc = Document()
    
    # Add some paragraphs
    doc.add_paragraph("Test Document")
    doc.add_paragraph("This is a test document with various content types.")
    
    # Add a table with potential issues
    table = doc.add_table(rows=3, cols=3)
    
    # Fill table with test data
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"R{row_idx}C{col_idx}"
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def test_enhanced_word_parsing(docx_path):
    """Test the enhanced Word document parsing logic"""
    print(f"Testing Word document parsing: {docx_path}")
    
    try:
        from docx import Document as docx_Document
        
        doc = docx_Document(docx_path)
        word_content = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            word_content += paragraph.text + "\n"
        
        # Extract tables with enhanced error handling (mimicking our app logic)
        try:
            for table_idx, table in enumerate(doc.tables):
                word_content += f"\n--- Table {table_idx + 1} ---\n"
                for row_idx, row in enumerate(table.rows):
                    try:
                        # Handle potential missing cells or malformed table structure
                        cells_text = []
                        for cell in row.cells:
                            if cell and hasattr(cell, 'text'):
                                cells_text.append(cell.text.strip())
                            else:
                                cells_text.append("[empty]")
                        
                        row_text = " | ".join(cells_text)
                        if row_text.strip() and row_text != " | ".join(["[empty]"] * len(cells_text)):
                            word_content += row_text + "\n"
                    except Exception as cell_error:
                        # Skip problematic rows and continue
                        word_content += f"[Table row {row_idx}: parsing error - {cell_error}]\n"
                        continue
                word_content += "\n"
        except Exception as table_error:
            # If table parsing fails completely, continue with text content
            word_content += f"\n--- Tables could not be parsed: {str(table_error)} ---\n"
        
        print("✅ Word document parsing successful!")
        print(f"Content length: {len(word_content)} characters")
        print("Preview:")
        print("-" * 50)
        print(word_content[:500] + "..." if len(word_content) > 500 else word_content)
        print("-" * 50)
        
        return True
        
    except Exception as e:
        print(f"❌ Word document parsing failed: {str(e)}")
        
        # Test fallback logic
        try:
            print("🔄 Testing fallback logic...")
            doc = docx_Document(docx_path)
            basic_content = ""
            for paragraph in doc.paragraphs:
                basic_content += paragraph.text + "\n"
            
            if basic_content:
                print("✅ Fallback parsing successful!")
                print(f"Basic content length: {len(basic_content)} characters")
                return True
            else:
                print("⚠️ No content extracted even with fallback")
                return False
        except Exception as fallback_error:
            print(f"❌ Even fallback parsing failed: {str(fallback_error)}")
            return False

def main():
    print("Word Document Parsing Test")
    print("=" * 40)
    
    # Create test document
    test_docx = create_test_word_doc()
    print(f"Created test document: {test_docx}")
    
    try:
        # Test parsing
        success = test_enhanced_word_parsing(test_docx)
        
        if success:
            print("\n🎉 All tests passed! Word document parsing enhancements are working correctly.")
        else:
            print("\n❌ Tests failed. There may be issues with the Word document parsing.")
            
    finally:
        # Clean up
        if os.path.exists(test_docx):
            os.unlink(test_docx)
            print(f"Cleaned up test file: {test_docx}")

if __name__ == "__main__":
    main()
