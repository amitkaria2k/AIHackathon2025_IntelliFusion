"""
Document Generator Service for creating documents from templates
"""
import os
from pathlib import Path
from typing import Dict, List, Optional
import jinja2
from docx import Document
from docx.shared import Inches
import json
from datetime import datetime
import hashlib

from ..config import Config
from ..database import get_session_local, Document as DocumentModel, DocumentRevision

class DocumentGeneratorService:
    """Service for generating documents from templates"""
    
    def __init__(self):
        """Initialize the document generator service"""
        self.templates_dir = Config.TEMPLATES_DIR
        self.documents_dir = Config.DOCUMENTS_DIR
        
        # Ensure directories exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize Jinja2 environment
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(self.templates_dir)),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
    
    def generate_document(self, 
                         project_id: str,
                         document_type: str, 
                         content: str, 
                         template_name: Optional[str] = None,
                         variables: Optional[Dict] = None) -> str:
        """
        Generate a document from content and template
        
        Args:
            project_id: ID of the project
            document_type: Type of document
            content: Generated content from LLM
            template_name: Name of template to use
            variables: Template variables
            
        Returns:
            Path to generated document
        """
        
        # Create document filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_name = f"{document_type.replace(' ', '_')}_{timestamp}"
        
        if template_name and self._template_exists(template_name):
            # Use Jinja2 template
            file_path = self._generate_from_template(
                doc_name, template_name, content, variables or {}
            )
        else:
            # Generate Word document directly
            file_path = self._generate_word_document(doc_name, document_type, content)
        
        # Save to database
        self._save_document_record(project_id, doc_name, document_type, file_path, template_name)
        
        return file_path
    
    def _template_exists(self, template_name: str) -> bool:
        """Check if template file exists"""
        template_path = self.templates_dir / f"{template_name}.html"
        return template_path.exists()
    
    def _generate_from_template(self, 
                              doc_name: str, 
                              template_name: str, 
                              content: str, 
                              variables: Dict) -> str:
        """Generate document from Jinja2 template"""
        
        try:
            template = self.jinja_env.get_template(f"{template_name}.html")
            
            # Prepare template variables
            template_vars = {
                'content': content,
                'generated_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'document_name': doc_name.replace('_', ' '),
                **variables
            }
            
            # Render template
            rendered_content = template.render(**template_vars)
            
            # Save as HTML file
            output_path = self.documents_dir / f"{doc_name}.html"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(rendered_content)
            
            return str(output_path)
            
        except Exception as e:
            raise Exception(f"Error generating document from template: {str(e)}")
    
    def _generate_word_document(self, doc_name: str, document_type: str, content: str) -> str:
        """Generate Word document directly"""
        
        try:
            # Create new Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(document_type, 0)
            title.alignment = 1  # Center alignment
            
            # Add metadata
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")  # Empty line
            
            # Split content into sections and add to document
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    if section.startswith('#'):
                        # Header
                        level = section.count('#')
                        header_text = section.lstrip('#').strip()
                        doc.add_heading(header_text, level)
                    elif section.strip().startswith('-') or section.strip().startswith('•'):
                        # List items
                        lines = section.split('\n')
                        for line in lines:
                            if line.strip():
                                p = doc.add_paragraph()
                                p.style = 'List Bullet'
                                p.add_run(line.lstrip('-•').strip())
                    else:
                        # Regular paragraph
                        doc.add_paragraph(section.strip())
            
            # Save document
            output_path = self.documents_dir / f"{doc_name}.docx"
            doc.save(str(output_path))
            
            return str(output_path)
            
        except Exception as e:
            raise Exception(f"Error generating Word document: {str(e)}")
    
    def _save_document_record(self, 
                            project_id: str, 
                            doc_name: str, 
                            document_type: str, 
                            file_path: str, 
                            template_name: Optional[str] = None):
        """Save document record to database"""
        
        try:
            SessionLocal = get_session_local()
            db = SessionLocal()
            
            # Calculate content hash for change detection
            content_hash = self._calculate_file_hash(file_path)
            
            # Create document record
            document = DocumentModel(
                project_id=project_id,
                name=doc_name,
                document_type=document_type,
                template_name=template_name,
                file_path=file_path,
                content_hash=content_hash,
                status="Draft"
            )
            
            db.add(document)
            db.commit()
            
            # Create initial revision
            revision = DocumentRevision(
                document_id=document.id,
                version="1.0",
                changes_description="Initial document creation",
                file_path=file_path,
                created_by="System"
            )
            
            db.add(revision)
            db.commit()
            db.close()
            
        except Exception as e:
            print(f"Error saving document record: {str(e)}")
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file content"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except:
            return ""
    
    def create_default_templates(self):
        """Create default document templates"""
        
        templates = {
            "project_management_plan": {
                "title": "Project Management Plan",
                "content": """
<html>
<head>
    <title>{{ document_name }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 2cm; }
        h1 { color: #00629B; border-bottom: 2px solid #00629B; }
        h2 { color: #00629B; }
        .meta { color: #666; font-size: 0.9em; }
    </style>
</head>
<body>
    <h1>{{ document_name }}</h1>
    <p class="meta">Generated on: {{ generated_date }}</p>
    
    <div class="content">
        {{ content | safe }}
    </div>
</body>
</html>
                """
            },
            "technical_concept": {
                "title": "Technical Concept Document",
                "content": """
<html>
<head>
    <title>{{ document_name }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 2cm; }
        h1 { color: #00629B; border-bottom: 2px solid #00629B; }
        h2 { color: #00629B; }
        .meta { color: #666; font-size: 0.9em; }
        .section { margin-bottom: 2em; }
    </style>
</head>
<body>
    <h1>{{ document_name }}</h1>
    <p class="meta">Generated on: {{ generated_date }}</p>
    
    <div class="content">
        {{ content | safe }}
    </div>
</body>
</html>
                """
            }
        }
        
        for template_name, template_data in templates.items():
            template_path = self.templates_dir / f"{template_name}.html"
            if not template_path.exists():
                with open(template_path, 'w', encoding='utf-8') as f:
                    f.write(template_data["content"])
